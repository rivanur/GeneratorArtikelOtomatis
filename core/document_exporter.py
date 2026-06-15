import os
import uuid
import requests
import io
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from bs4 import BeautifulSoup
from markdown_it import MarkdownIt

class DocumentExporter:
    @staticmethod
    def export_to_word(title: str, markdown_content: str, output_dir: str = "data/outputs") -> str:
        """
        Mengonversi teks Markdown menjadi file Word (.docx)
        Mengembalikan path file (string)
        """
        os.makedirs(output_dir, exist_ok=True)
        
        # 1. Konversi Markdown ke HTML
        md = MarkdownIt()
        # Buang baris H1 pertama jika ada, karena kita akan buat title_para sendiri
        lines = markdown_content.split('\n')
        if lines and lines[0].startswith('# '):
            lines = lines[1:]
        cleaned_markdown = '\n'.join(lines).strip()
        html_content = md.render(cleaned_markdown)

        # 2. Buat Dokumen Word Baru
        doc = Document()
        
        # 3. Tambahkan Judul Utama
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_para.alignment = 1 # Center

        doc.add_paragraph() # Spasi

        # 4. Parse HTML menggunakan BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'li']):
            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'p':
                img = element.find('img')
                if img and img.get('src'):
                    # Ada gambar, jalankan "Penyedot Gambar"
                    img_url = img.get('src')
                    try:
                        response = requests.get(img_url, timeout=15)
                        if response.status_code == 200:
                            image_bytes = response.content
                            try:
                                # Konversi gambar ke PNG menggunakan Pillow (hindari error format webp)
                                img_pil = Image.open(io.BytesIO(image_bytes))
                                if img_pil.mode in ("RGBA", "P"):
                                    img_pil = img_pil.convert("RGB")
                                
                                converted_stream = io.BytesIO()
                                img_pil.save(converted_stream, format="PNG")
                                converted_stream.seek(0)
                                image_stream = converted_stream
                            except Exception as pil_err:
                                print(f"PIL warning: {pil_err}")
                                image_stream = io.BytesIO(image_bytes)

                            # Lebar di-set max 6 inci agar rapi di kertas A4
                            doc.add_picture(image_stream, width=Inches(6.0))
                            
                            alt_text = img.get('alt')
                            if alt_text:
                                caption_para = doc.add_paragraph(alt_text)
                                caption_para.alignment = 1 # Center
                                caption_para.runs[0].font.italic = True
                                caption_para.runs[0].font.size = Pt(9)
                        else:
                            doc.add_paragraph(f"[Gambar Gagal Dimuat: {img_url}]")
                    except Exception as e:
                        print(f"Gagal mendownload gambar {img_url}: {e}")
                        doc.add_paragraph(f"[Tautan Gambar Tidak Valid: {img_url}]")
                else:
                    doc.add_paragraph(element.get_text())
            elif element.name == 'li':
                # Sederhana: gunakan style bawaan List Bullet
                parent = element.find_parent()
                style = 'List Number' if parent and parent.name == 'ol' else 'List Bullet'
                try:
                    doc.add_paragraph(element.get_text(), style=style)
                except KeyError:
                    # Fallback jika style default tidak tersedia di python-docx
                    prefix = "- " if style == 'List Bullet' else "1. "
                    doc.add_paragraph(prefix + element.get_text())

        # 5. Simpan ke File
        safe_title = "".join([c if c.isalnum() else "_" for c in title])[:30]
        filename = f"{safe_title}_{uuid.uuid4().hex[:6]}.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc.save(filepath)
        return filepath
